import os
import sys
from docx import Document
#from docx.shared import Pt, RGBColor
#from docx.oxml.ns import qn


# Assuming the translation_agent module is in the src/translation_agent directory
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import src.translation_agent.utils as ta

def read_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # Check if the paragraph is not empty
                full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Error reading .docx file: {e}")

def convert_doc_to_docx(file_path):
    try:
        from win32com.client import Dispatch
        word = Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(file_path)
        docx_file_path = os.path.splitext(file_path)[0] + ".docx"
        doc.SaveAs(docx_file_path, FileFormat=12)  # 12 represents the wdFormatXMLDocument
        doc.Close()
        word.Quit()
        return docx_file_path
    except ImportError:
        raise ImportError("Conversion of .doc to .docx files is only supported on Windows with pywin32 installed.")
    except Exception as e:
        raise Exception(f"Error converting .doc to .docx: {e}")

def read_word_file(file_path):
    try:
        file_path_lower = file_path.lower()
        if file_path_lower.endswith('.docx'):
            return read_docx(file_path)
        elif file_path_lower.endswith('.doc'):
            if sys.platform == "win32":
                # Convert .doc to .docx then read it
                docx_file_path = convert_doc_to_docx(file_path)
                return read_docx(docx_file_path)
            else:
                raise ValueError("Reading .doc files is only supported on Windows.")
        else:
            raise ValueError("Unsupported file format. Please provide a .docx or .doc file.")
    except Exception as e:
        raise Exception(f"Error processing the Word file: {e}")

def write_translated_docx(file_path, translated_text):
    # Generate new file path
    new_file_path = file_path.rsplit('.', 1)
    new_file_path = f"{new_file_path[0]}_translated.docx"

    # Create a new Document
    doc = Document()

    # Split text by newlines into paragraphs
    paragraphs = translated_text.split('\n')

    # Exclude the first and last paragraphs and add the rest to the document
    #for line in paragraphs[1:-1]:
    for line in paragraphs:
        doc.add_paragraph(line)

    # Save the document
    doc.save(new_file_path)    
    
    return new_file_path

def remove_paragraphs_containing_text(input_file_path, output_file_path):
    # Load the document
    doc = Document(input_file_path)
    # Find and remove paragraphs containing the specified text
    paragraphs_to_remove = [p for p in doc.paragraphs if ("TRANSLATION" in p.text or "TRANSLATE" in p.text) ]
    for paragraph in paragraphs_to_remove:
        p = paragraph._element
        p.getparent().remove(p)
        p._element = p._p = None
    
    # Save the modified document
    doc.save(output_file_path)

def preserve_format_and_replace_text(original_file_path, translated_file_path, output_file_path):
    # Load the original and translated documents
    original_doc = Document(original_file_path)
    translated_doc = Document(translated_file_path)
    
    # Prepare to iterate over paragraphs in the original and translated documents
    original_paragraphs = original_doc.paragraphs
    translated_paragraphs = translated_doc.paragraphs
    
    # Index for the translated paragraphs
    translated_index = 0
    
    # Loop through original paragraphs and replace with corresponding translated paragraphs
    for i in range(len(original_paragraphs)):
        original_paragraph = original_paragraphs[i]

        # Skip the original paragraph if it's empty
        if not original_paragraph.text.strip():
            continue
        
        # Find the next non-empty translated paragraph
        while translated_index < len(translated_paragraphs) and not translated_paragraphs[translated_index].text.strip():
            translated_index += 1
        
        # If there are no more translated paragraphs, break the loop
        if translated_index >= len(translated_paragraphs):
            break
        
        translated_paragraph = translated_paragraphs[translated_index]
        
        # Replace the text in the original paragraph with text from the translated paragraph
        #original_paragraph.add_run(translated_paragraph.text)
        if original_paragraph.runs:
            first_run_format = original_paragraph.runs[0]
            original_paragraph.clear()
        
            # Add the translated text with the same formatting
            new_run = original_paragraph.add_run(translated_paragraph.text)
            new_run.bold = first_run_format.bold
            new_run.italic = first_run_format.italic
            new_run.underline = first_run_format.underline
            new_run.font.size = first_run_format.font.size
            new_run.font.name = first_run_format.font.name
        else:
            # If the original paragraph has no runs, simply add the text from the translated paragraph
            original_paragraph.clear()
            original_paragraph.add_run(translated_paragraph.text)
        
        # Move to the next translated paragraph
        translated_index += 1
    
    # Save the modified document to the output file path
    original_doc.save(output_file_path)

def translate_table(input_file_path, output_file_path):
    # Load the Word document
    doc = Document(input_file_path)
    
    # Iterate through all tables in the document
    for table in doc.tables:
        # Initialize an empty list to hold the cell contents
        original_table_content = []
        
        # Collect the contents of the table cells
        for row in table.rows:
            for cell in row.cells:
                cell_content = '\n'.join([paragraph.text for paragraph in cell.paragraphs])
                original_table_content.append(cell_content)
        
        # Join the cell contents into a single string, separating by new lines
        entire_table_content = '\n'.join(original_table_content)
        
        # Translate the entire table content
        source_lang, target_lang, country = "English", "Chinese", "China"
        translated_text = ta.translate(
            source_lang=source_lang,
            target_lang=target_lang,
            source_text=entire_table_content,
            country=country,
        )
        
        # Split the translated content back into individual cell contents
        translated_contents = translated_text.split('\n')
        translated_contents = [text for text in translated_contents if "TRANSLATION" not in text and "TRANSLATE" not in text]

        # Replace the original cell contents with the translated ones
        index = 0
        for row in table.rows:
            for cell in row.cells:
                original_paragraphs = cell.paragraphs
                if index < len(translated_contents):
                    translated_text = translated_contents[index]
                    
                    for original_paragraph in original_paragraphs:
                        if original_paragraph.runs:
                            first_run_format = original_paragraph.runs[0]
                            original_paragraph.clear()
                            
                            # Add the translated text with the same formatting
                            new_run = original_paragraph.add_run(translated_text)
                            new_run.bold = first_run_format.bold
                            new_run.italic = first_run_format.italic
                            new_run.underline = first_run_format.underline
                            new_run.font.size = first_run_format.font.size
                            new_run.font.name = first_run_format.font.name
                        else:
                            # If the original paragraph has no runs, simply add the translated text
                            original_paragraph.clear()
                            original_paragraph.add_run(translated_text)
                index += 1

    # Save the translated document to the output path
    doc.save(output_file_path)


if __name__ == "__main__":
    source_lang, target_lang, country = "English", "Chinese", "China"

    file_path = input("Enter the path to the Word document: ")
    if not os.path.isfile(file_path):
        print("The specified file does not exist.")
        sys.exit(1)

    try:
        source_text = read_word_file(file_path)
        print(f"Source text:\n\n{source_text}\n------------\n")

        translation = ta.translate(
            source_lang=source_lang,
            target_lang=target_lang,
            source_text=source_text,
            country=country,
        )

        #print(f"Translation:\n\n{translation}")
        translated_file = write_translated_docx(file_path, translation)

        #remove strings like <TRANSLATION>
        #remove_paragraphs_containing_text(translated_file, translated_file)

        #preserve_format_and_replace_text
        #preserve_format_and_replace_text(file_path, translated_file, translated_file)

        #process table translation
        #translate_table(translated_file, translated_file)

        print(f"The translated document has been saved as: {translated_file}")
        
    except Exception as e:
        print(f"An error occurred: {e}")
