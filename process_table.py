import os
import sys
from docx import Document
from markdownify import markdownify as md
from docx.table import Table
import re

# Assuming the translation_agent module is in the src/translation_agent directory
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import src.translation_agent.utils as ta

def read_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Error reading .docx file: {e}")

def convert_doc_to_docx(file_path):
    try:
        from win32com.client import Dispatch
        word = Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(file_path)
        docx_file_path = os.path.splitext(file_path)[0] + ".docx"
        doc.SaveAs(docx_file_path, FileFormat=12)  # 12 represents the wdFormatXMLDocument
        doc.Close()
        word.Quit()
        return docx_file_path
    except ImportError:
        raise ImportError("Conversion of .doc to .docx files is only supported on Windows with pywin32 installed.")
    except Exception as e:
        raise Exception(f"Error converting .doc to .docx: {e}")

def read_word_file(file_path):
    try:
        file_path_lower = file_path.lower()
        if file_path_lower.endswith('.docx'):
            return read_docx(file_path)
        elif file_path_lower.endswith('.doc'):
            if sys.platform == "win32":
                # Convert .doc to .docx then read it
                docx_file_path = convert_doc_to_docx(file_path)
                return read_docx(docx_file_path)
            else:
                raise ValueError("Reading .doc files is only supported on Windows.")
        else:
            raise ValueError("Unsupported file format. Please provide a .docx or .doc file.")
    except Exception as e:
        raise Exception(f"Error processing the Word file: {e}")

def write_translated_docx(file_path, translated_text):
    # Generate new file path
    new_file_path = file_path.rsplit('.', 1)
    new_file_path = f"{new_file_path[0]}_translated.docx"

    # Create a new Document
    doc = Document()

    # Split text by newlines into paragraphs
    paragraphs = translated_text.split('\n')

    # Exclude the first and last paragraphs and add the rest to the document
    #for line in paragraphs[1:-1]:
    for line in paragraphs:

        doc.add_paragraph(line)

    # Save the document
    doc.save(new_file_path)    
    
    return new_file_path

# def translate_table(input_file_path, output_file_path):
#     # Load the Word document
#     doc = Document(input_file_path)
    
#     # Iterate through all tables in the document
#     for table in doc.tables:
#         # Initialize an empty list to hold the cell contents
#         original_table_content = []
        
#         # Collect the contents of the table cells
#         for row in table.rows:
#             for cell in row.cells:
#                 cell_content = '\n'.join([paragraph.text for paragraph in cell.paragraphs])
#                 original_table_content.append(cell_content)
        
#         # Join the cell contents into a single string, separating by new lines
#         entire_table_content = '\n'.join(original_table_content)
        
#         # Translate the entire table content
#         source_lang, target_lang, country = "English", "Chinese", "China"
#         translated_table_content = ta.translate(
#             source_lang=source_lang,
#             target_lang=target_lang,
#             source_text=entire_table_content,
#             country=country,
#         )
        
#         # Split the translated content back into individual cell contents
#         translated_contents = translated_table_content.split('\n')
#         translated_contents = [text for text in translated_contents if "TRANSLATION" not in text and "TRANSLATE" not in text]

        
#         # Replace the original cell contents with the translated ones
#         index = 0
#         for row in table.rows:
#             for cell in row.cells:
#                 original_paragraphs = cell.paragraphs
#                 if index < len(translated_contents):
#                     translated_text = translated_contents[index]
                    
#                     for original_paragraph in original_paragraphs:
#                         if original_paragraph.runs:
#                             first_run_format = original_paragraph.runs[0]
#                             original_paragraph.clear()
                            
#                             # Add the translated text with the same formatting
#                             new_run = original_paragraph.add_run(translated_text)
#                             new_run.bold = first_run_format.bold
#                             new_run.italic = first_run_format.italic
#                             new_run.underline = first_run_format.underline
#                             new_run.font.size = first_run_format.font.size
#                             new_run.font.name = first_run_format.font.name
#                         else:
#                             # If the original paragraph has no runs, simply add the translated text
#                             original_paragraph.clear()
#                             original_paragraph.add_run(translated_text)
#                 index += 1

#     # Save the translated document to the output path
#     doc.save(output_file_path)

def copy_run_format(source_run, target_run):
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb
    target_run.font.all_caps = source_run.font.all_caps
    target_run.font.small_caps = source_run.font.small_caps
    target_run.font.strike = source_run.font.strike
    
def translate_table(input_file_path, output_file_path):
    # Load the input document
    doc = Document(input_file_path)
    
    # Loop through every table in the document
    for table in doc.tables:
        # Loop through every row in the table
        for row in table.rows:
            # Loop through every cell in the row
            for cell in row.cells:
                if cell.text.strip():  # Skip empty or whitespace-only cells
                    if re.search(r'\w', cell.text):  # Check if cell contains words
                        source_lang, target_lang, country = "English", "Chinese", "China"
                        translated_text = ta.translate(
                            source_lang=source_lang,
                            target_lang=target_lang,
                            source_text=cell.text,
                            country=country,
                        )

                        translated_contents = translated_text.split('\n')
                        translated_contents = [text for text in translated_contents if "TRANSLATION" not in text and "TRANSLATE" not in text]
                        translated_text = '\n'.join(translated_contents)

                    # Directly replace the cell text with the translated text
                    cell.text = translated_text

    doc.save(output_file_path)


if __name__ == "__main__":

    file_path = input("Enter the path to the Word document: ")
    if not os.path.isfile(file_path):
        print("The specified file does not exist.")
        sys.exit(1)

    try:
        new_file_path = file_path.rsplit('.', 1)
        new_file_path = f"{new_file_path[0]}_translated.docx"

        translate_table(file_path, new_file_path)
        
        #print(f"The translated document has been saved as: {new_file_path}")
        
    except Exception as e:
        print(f"An error occurred: {e}")
