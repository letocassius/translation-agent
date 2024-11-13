import os
import sys
from docx import Document
#from docx.shared import Pt, RGBColor
#from docx.oxml.ns import qn


# Assuming the translation_agent module is in the src/translation_agent directory
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import src.translation_agent.utils as ta

def read_word_file(file_path):
    doc = Document(file_path)
    text_paragraphs = [para for para in doc.paragraphs if para.text.strip()]
    return text_paragraphs

def translate_table(input_file_path, output_file_path):
    # Load the Word document
    doc = Document(input_file_path)
    
    # Iterate through all tables in the document
    for table in doc.tables:
        # Initialize an empty list to hold the cell contents
        original_table_content = []
        
        # Collect the contents of the table cells
        for row in table.rows:
            for cell in row.cells:
                #cell_content = '\n'.join([paragraph.text for paragraph in cell.paragraphs])
                cell_content = cell.text
                original_table_content.append(cell_content)  
        
        # Join the cell contents into a single string, separating by new lines
        entire_table_content = '\n'.join(original_table_content)
        
        # Translate the entire table content
        source_lang, target_lang, country = "English", "Chinese", "China"
        translated_table_content = ta.translate(
            source_lang=source_lang,
            target_lang=target_lang,
            source_text=entire_table_content,
            country=country,
        )
        
        # Split the translated content back into individual cell contents
        translated_contents = translated_table_content.split('\n')
        translated_contents = [text for text in translated_contents if "TRANSLATION" not in text and "TRANSLATE" not in text]
        
        # Replace the original cell contents with the translated ones
        index = 0
        for row in table.rows:
            for cell in row.cells:
                if index < len(translated_contents):
                    translated_text = translated_contents[index]
                    
                    cell.text = translated_text
                    index += 1
                else:
                    break

    # Save the translated document to the output path
    doc.save(output_file_path)

def preserve_format_and_replace_text(file_path, translated_paragraphs, start_paragraph_num, chunk_size):
    # Open the original document
    doc = Document(file_path)
    
    original_paragraphs = doc.paragraphs
    translated_index = 0  # Index for the translated paragraphs array

    # Define the end paragraph number for the chunk
    end_paragraph_num = start_paragraph_num + chunk_size

    for i in range(start_paragraph_num, min(end_paragraph_num, len(original_paragraphs))):
        original_paragraph = original_paragraphs[i]

        # Skip the original paragraph if it's empty
        if not original_paragraph.text.strip():
            continue

        # Find the next non-empty translated paragraph
        while (translated_index < len(translated_paragraphs) and not translated_paragraphs[translated_index].strip()):
            translated_index += 1
        
        # If there are no more translated paragraphs, break the loop
        if translated_index >= len(translated_paragraphs):
            break
        
        translated_paragraph_text = translated_paragraphs[translated_index]

        # If the original paragraph has no runs, we need to add a dummy run to capture formatting
        if not original_paragraph.runs:
            original_paragraph.add_run("")

        # Extract formatting from the first run of the original paragraph
        first_run_format = original_paragraph.runs[0]

        # Clear the original paragraph's text while keeping the existing runs
        for run in original_paragraph.runs:
            run.text = ""

        # Add new translated text into the original paragraph while preserving formatting
        for run in translated_paragraph_text.split('\n'):
            new_run = original_paragraph.add_run(run)
            new_run.bold = first_run_format.bold
            new_run.italic = first_run_format.italic
            new_run.underline = first_run_format.underline
            if first_run_format.font.size:
                new_run.font.size = first_run_format.font.size
            if first_run_format.font.name:
                new_run.font.name = first_run_format.font.name

        # Move to the next translated paragraph
        translated_index += 1

    # Save the modified document
    doc.save(file_path)

if __name__ == "__main__":
    original_file_path = input("Enter the path to the Word document: ")

    doc = Document(original_file_path)
    file_path = original_file_path.rsplit('.', 1)
    output_file_path = f"{file_path[0]}_translated.docx"
    doc.save(output_file_path)
    
    if not os.path.isfile(original_file_path):
        print("The specified file does not exist.")
        sys.exit(1)

    try:
        paragraphs = read_word_file(original_file_path)
        total_paragraphs = len(paragraphs)
        chunk_size = 50     # number of paragraphs per translation

        for i in range(0, total_paragraphs, chunk_size):
            translated_paragraphs = []
            chunk = [para.text for para in paragraphs[i:min(i + chunk_size,len(paragraphs))]]
            source_text = "\n".join(chunk)
            print(f"Processing paragraphs {i} to {i+chunk_size}...")

            source_lang, target_lang, country = "English", "Chinese", "China"
            translation = ta.translate(
                source_lang=source_lang,
                target_lang=target_lang,
                source_text=source_text,
                country=country
            )
            # Add the translated text to the list and filter out unwanted strings
            translated_paragraphs.extend(
                para for para in translation.split("\n") if "TRANSLATION" not in para and "TRANSLATE" not in para 
            )

            # replace the original text with translated text from line i to i+chunk_size
            preserve_format_and_replace_text(output_file_path, translated_paragraphs, i, chunk_size)

        # Translate tables in the document
        translate_table(output_file_path, output_file_path)

        print(f"The translated document has been saved as: {output_file_path}")

    except Exception as e:
        print(f"An error occurred: {e}")