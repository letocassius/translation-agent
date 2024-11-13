import streamlit as st
import os, sys
import re
import zipfile
from io import BytesIO
from docx import Document
from tempfile import NamedTemporaryFile, TemporaryDirectory

# Ensure the `src` directory is in the Python path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import src.translation_agent.utils as ta

def read_word_file(file_path):
    doc = Document(file_path)
    text_paragraphs = [para for para in doc.paragraphs if para.text.strip()]
    return text_paragraphs

def translate_table(input_file_path, output_file_path, source_lang, target_lang, country):
    doc = Document(input_file_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    if re.search(r'\w', cell.text):
                        translated_text = ta.translate(
                            source_lang=source_lang,
                            target_lang=target_lang,
                            source_text=cell.text,
                            country=country,
                        )
                        translated_contents = translated_text.split('\n')
                        translated_contents = [text for text in translated_contents if "TRANSLATION" not in text and "TRANSLATE" not in text]
                        translated_text = '\n'.join(translated_contents)
                    cell.text = translated_text
    doc.save(output_file_path)

def preserve_format_and_replace_text(file_path, translated_paragraphs, start_paragraph_num, chunk_size):
    doc = Document(file_path)
    original_paragraphs = doc.paragraphs
    translated_index = 0
    end_paragraph_num = start_paragraph_num + chunk_size
    for i in range(start_paragraph_num, min(end_paragraph_num, len(original_paragraphs))):
        original_paragraph = original_paragraphs[i]
        if not original_paragraph.text.strip():
            continue
        while (translated_index < len(translated_paragraphs) and not translated_paragraphs[translated_index].strip()):
            translated_index += 1
        if translated_index >= len(translated_paragraphs):
            break
        translated_paragraph_text = translated_paragraphs[translated_index]
        if not original_paragraph.runs:
            original_paragraph.add_run("")
        first_run_format = original_paragraph.runs[0]
        for run in original_paragraph.runs:
            run.text = ""
        for run in translated_paragraph_text.split('\n'):
            new_run = original_paragraph.add_run(run)
            new_run.bold = first_run_format.bold
            new_run.italic = first_run_format.italic
            new_run.underline = first_run_format.underline
            if first_run_format.font.size:
                new_run.font.size = first_run_format.font.size
            if first_run_format.font.name:
                new_run.font.name = first_run_format.font.name
        translated_index += 1
    doc.save(file_path)

def process_file(input_file_path, output_file_path, source_lang, target_lang, country, progress_callback=None):
    doc = Document(input_file_path)
    doc.save(output_file_path)
    try:
        paragraphs = read_word_file(input_file_path)
        total_paragraphs = len(paragraphs)
        chunk_size = 25    #number of paragraph per translation 
        for i in range(0, total_paragraphs, chunk_size):
            translated_paragraphs = []
            chunk = [para.text for para in paragraphs[i:min(i + chunk_size, total_paragraphs)]]
            source_text = "\n".join(chunk)
            translation = ta.translate(
                source_lang=source_lang,
                target_lang=target_lang,
                source_text=source_text,
                country=country
            )
            translated_paragraphs.extend(
                para for para in translation.split("\n") if "TRANSLATION" not in para and "TRANSLATE" not in para
            )
            preserve_format_and_replace_text(output_file_path, translated_paragraphs, i, chunk_size)
            # Update the progress bar
            if progress_callback:
                progress_callback(i + chunk_size, total_paragraphs)
        translate_table(output_file_path, output_file_path, source_lang, target_lang, country)
    except Exception as e:
        print(f"An error occurred while processing {input_file_path}: {e}")

# The translation app
st.title("Translate Word File")
st.write("**Note: Translating one document might take about 3-5 minutes depending of the length of the document. Do not close the tab during the translation process or the results will not be saved.**")
st.write("**The default is set for uploading one file at a time. If you need to translate multiple files, make change to the env file.\
Translating multiple files requires longer time, do not close the tab during the translation process or the results will not be saved.**")

# Read the MULTIPLE_FILES_UPLOAD parameter from environment variable
accept_multiple_files = os.getenv("MULTIPLE_FILES_UPLOAD", "false").lower() == "true"

# Determine the file uploader based on MULTIPLE_FILES_UPLOAD value
if accept_multiple_files:
    uploaded_files = st.file_uploader("Upload Word Documents", type=["docx"], accept_multiple_files=True)
else:
    uploaded_files = [st.file_uploader("Upload a Word Document", type=["docx"])]

# Add radio buttons for translation direction
translate_option = st.radio(
    "Select translation direction:",
    ("English to Chinese", "Chinese to English")
)

# Determine the source and target languages based on user selection
if translate_option == "English to Chinese":
    source_lang, target_lang, country = "English", "Chinese", "China"
else:
    source_lang, target_lang, country = "Chinese", "English", "United States"

# Temporary directory to store translated files
translated_files = []

if st.button("Start Translate") and uploaded_files:
    with TemporaryDirectory() as temp_dir:
        for uploaded_file in uploaded_files:
            with NamedTemporaryFile(delete=False, suffix=".docx") as temp_input_file:
                temp_input_file.write(uploaded_file.read())
                temp_input_file_path = temp_input_file.name

            input_file_name = uploaded_file.name
            output_file_name = f"{os.path.splitext(input_file_name)[0]}_translated.docx"
            temp_output_file_path = os.path.join(temp_dir, output_file_name)

            st.write(f"Translating document: {input_file_name}")

            # Initialize progress bar and text
            progress_bar = st.progress(0)
            progress_text = st.empty()

            def update_progress(current, total):
                progress = min((current / total),0.9)
                progress_bar.progress(progress)
                progress_text.text(f"Translation progress... ({int(progress * 100)}%)")

            process_file(temp_input_file_path, temp_output_file_path, source_lang, target_lang, country, update_progress)

            # Mark progress as complete and show completion message
            progress_bar.progress(1.0)
            progress_text.text("Translation has completed.")

            translated_files.append(temp_output_file_path)
        
        # download translated file either individually or in a zip file
        if accept_multiple_files:
            # Create a Zip file for all translated files
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zf:
                for translated_file in translated_files:
                    zf.write(translated_file, os.path.basename(translated_file))

            zip_buffer.seek(0)
            # Provide a button to download the zip file
            st.download_button("Download All Translated Documents", data=zip_buffer, file_name="translated_documents.zip", mime="application/zip")
        else:
            with open(temp_output_file_path, "rb") as output_file:
                st.download_button(f"Download {output_file_name}", data=output_file.read(), file_name=output_file_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.write("Upload the files, choose the translation direction, and click 'Start Translate' to translate and download the documents.")
