import os
import sys
from docx import Document
import re

# Assuming the translation_agent module is in the src/translation_agent directory
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import src.translation_agent.utils as ta

def read_word_file(file_path):
    doc = Document(file_path)
    text_paragraphs = [para for para in doc.paragraphs if para.text.strip()]
    return text_paragraphs

def translate_table(input_file_path, output_file_path):
    # Load the input document
    doc = Document(input_file_path)
    
    # Loop through every table in the document
    for table in doc.tables:
        # Loop through every row in the table
        for row in table.rows:
            # Loop through every cell in the row
            for cell in row.cells:
                if cell.text.strip():  # Skip empty or whitespace-only cells
                    if re.search(r'\w', cell.text):  # Check if cell contains words
                        source_lang, target_lang, country = "English", "Chinese", "China"
                        translated_text = ta.translate(
                            source_lang=source_lang,
                            target_lang=target_lang,
                            source_text=cell.text,
                            country=country,
                        )

                        translated_contents = translated_text.split('\n')
                        translated_contents = [text for text in translated_contents if "TRANSLATION" not in text and "TRANSLATE" not in text]
                        translated_text = '\n'.join(translated_contents)

                    # Directly replace the cell text with the translated text
                    cell.text = translated_text

    doc.save(output_file_path)

def preserve_format_and_replace_text(file_path, translated_paragraphs, start_paragraph_num, chunk_size):
    # Open the original document
    doc = Document(file_path)
    
    original_paragraphs = doc.paragraphs
    translated_index = 0  # Index for the translated paragraphs array

    # Define the end paragraph number for the chunk
    end_paragraph_num = start_paragraph_num + chunk_size

    for i in range(start_paragraph_num, min(end_paragraph_num, len(original_paragraphs))):
        original_paragraph = original_paragraphs[i]

        # Skip the original paragraph if it's empty
        if not original_paragraph.text.strip():
            continue

        # Find the next non-empty translated paragraph
        while (translated_index < len(translated_paragraphs) and not translated_paragraphs[translated_index].strip()):
            translated_index += 1
        
        # If there are no more translated paragraphs, break the loop
        if translated_index >= len(translated_paragraphs):
            break
        
        translated_paragraph_text = translated_paragraphs[translated_index]

        # If the original paragraph has no runs, we need to add a dummy run to capture formatting
        if not original_paragraph.runs:
            original_paragraph.add_run("")

        # Extract formatting from the first run of the original paragraph
        first_run_format = original_paragraph.runs[0]

        # Clear the original paragraph's text while keeping the existing runs
        for run in original_paragraph.runs:
            run.text = ""

        # Add new translated text into the original paragraph while preserving formatting
        for run in translated_paragraph_text.split('\n'):
            new_run = original_paragraph.add_run(run)
            new_run.bold = first_run_format.bold
            new_run.italic = first_run_format.italic
            new_run.underline = first_run_format.underline
            if first_run_format.font.size:
                new_run.font.size = first_run_format.font.size
            if first_run_format.font.name:
                new_run.font.name = first_run_format.font.name

        # Move to the next translated paragraph
        translated_index += 1

    # Save the modified document
    doc.save(file_path)

def process_file(file_path, output_folder):
    doc = Document(file_path)
    file_name = os.path.basename(file_path).rsplit('.', 1)[0]
    output_file_path = os.path.join(output_folder, f"{file_name}_translated.docx")
    doc.save(output_file_path)
    
    try:
        paragraphs = read_word_file(file_path)  # Function should be defined elsewhere
        total_paragraphs = len(paragraphs)
        chunk_size = 25     # Number of paragraphs per translation

        for i in range(0, total_paragraphs, chunk_size):
            translated_paragraphs = []
            chunk = [para.text for para in paragraphs[i:min(i + chunk_size, total_paragraphs)]]
            source_text = "\n".join(chunk)
            print(f"Processing paragraphs {i} to {i + chunk_size}...")

            source_lang, target_lang, country = "English", "Chinese", "China"
            translation = ta.translate(
                source_lang=source_lang,
                target_lang=target_lang,
                source_text=source_text,
                country=country
            )
            # Add the translated text to the list and filter out unwanted strings
            translated_paragraphs.extend(
                para for para in translation.split("\n") if "TRANSLATION" not in para and "TRANSLATE" not in para
            )

            # Replace the original text with translated text from line i to i+chunk_size
            preserve_format_and_replace_text(output_file_path, translated_paragraphs, i, chunk_size)  # Function should be defined elsewhere

        # Translate tables in the document
        translate_table(output_file_path, output_file_path)  # Function should be defined elsewhere

        print(f"The translated document has been saved as: {output_file_path}")

    except Exception as e:
        print(f"An error occurred while processing {file_path}: {e}")

if __name__ == "__main__":
    input_folder_path = input("Enter the input folder path: ")
    output_folder_path = input("Enter the output folder path: ")

    if not os.path.isdir(input_folder_path):
        print("The specified input folder does not exist.")
        sys.exit(1)

    if not os.path.isdir(output_folder_path):
        print("The specified output folder does not exist.")
        sys.exit(1)

    try:
        for filename in os.listdir(input_folder_path):
            if filename.endswith(".docx"):
                file_path = os.path.join(input_folder_path, filename)
                process_file(file_path, output_folder_path)
        print("All files processed successfully.")
    except Exception as e:
        print(f"An error occurred: {e}")