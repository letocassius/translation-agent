from docx import Document
import os
import sys

# Add the path to the src.translation_agent.utils module
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import src.translation_agent.utils as ta

# Function to convert .doc to .docx
def convert_doc_to_docx(doc_path):
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(doc_path)
    new_path = doc_path + "x"
    doc.SaveAs(new_path, FileFormat=16)  # FileFormat=16 for .docx
    doc.Close()
    word.Quit()
    return new_path

# Load the document
file_path = input("Enter the path to the Word document: ")
if not os.path.isfile(file_path):
    print("The specified file does not exist.")
    sys.exit(1)

# Convert .doc to .docx if necessary
if file_path.lower().endswith('.doc'):
    file_path = convert_doc_to_docx(file_path)

source_text = Document(file_path)

count = 0
for paragraph in source_text.paragraphs:
    #print(f"Paragraph {count}: {paragraph.text}")
    count += 1

print(count)
    


# Example of how to loop through tables (uncomment if needed)
# print("\nTables:")
# for table in source_text.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text)